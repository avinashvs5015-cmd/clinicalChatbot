"""
Document extractor module for RAG system.
Supports extraction from PDF and DOC/DOCX files.
"""

import os
from typing import List, Dict, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats."""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.doc', '.docx', '.txt'}
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num} in {file_path}: {e}")
                        continue
            
            if not text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file (legacy format)."""
        # For DOC files, we'll need a different approach
        # This is a placeholder - in practice, you might use python-docx2txt or antiword
        logger.warning(f"DOC format not fully supported yet: {file_path}")
        return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                logger.warning(f"No text extracted from TXT: {file_path}")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from a document file based on its extension."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            logger.error(f"Unsupported file format: {extension}")
            return ""
        
        logger.info(f"Extracting text from: {file_path.name}")
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif extension == '.docx':
            return self.extract_text_from_docx(str(file_path))
        elif extension == '.doc':
            return self.extract_text_from_doc(str(file_path))
        elif extension == '.txt':
            return self.extract_text_from_txt(str(file_path))
        else:
            logger.error(f"Unsupported extension: {extension}")
            return ""
    
    def extract_from_directory(self, directory_path: str) -> Dict[str, str]:
        """Extract text from all supported documents in a directory."""
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return {}
        
        extracted_texts = {}
        
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                text = self.extract_text(str(file_path))
                if text:
                    extracted_texts[str(file_path)] = text
                    logger.info(f"Successfully extracted {len(text)} characters from {file_path.name}")
                else:
                    logger.warning(f"No text extracted from {file_path.name}")
        
        logger.info(f"Extracted text from {len(extracted_texts)} documents")
        return extracted_texts
    
    def get_document_metadata(self, file_path: str) -> Dict[str, any]:
        """Get metadata for a document."""
        file_path = Path(file_path)
        
        metadata = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'extension': file_path.suffix.lower(),
            'file_size': file_path.stat().st_size if file_path.exists() else 0,
            'modified_time': file_path.stat().st_mtime if file_path.exists() else 0
        }
        
        return metadata


# Example usage
if __name__ == "__main__":
    extractor = DocumentExtractor()
    
    # Test with a directory
    documents_dir = "../documents"
    extracted_texts = extractor.extract_from_directory(documents_dir)
    
    for file_path, text in extracted_texts.items():
        print(f"\n--- {Path(file_path).name} ---")
        print(f"Text length: {len(text)} characters")
        print(f"Preview: {text[:200]}...")